from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("Local_Network_Messenger_in_Go.docx")
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT = "F5F7F9"
BORDER = "DADCE0"

def set_font(run, name="Arial", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)

def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margin(cell)

def add_p(doc, text="", after=8, before=0, size=11, bold=False, italic=False, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    cell_margin(cell, 100, 140, 100, 140)
    cell.text = ""
    for i, line in enumerate(code.strip("\n").splitlines()):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line or " ")
        set_font(r, name="Courier New", size=8.5)
    add_p(doc, "", after=3)

def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    cell_margin(cell, 100, 140, 100, 140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    set_font(r, size=10, bold=True)
    r = p.add_run(text)
    set_font(r, size=10)
    add_p(doc, "", after=2)

def add_bullet(doc, text, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size=11)
    return p

def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size=11)
    return p

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"; normal.font.size = Pt(11); normal.font.color.rgb = BLACK
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.15
for name, size, color, before, after in [
    ("Heading 1", 20, BLACK, 20, 6), ("Heading 2", 16, BLACK, 18, 6), ("Heading 3", 14, RGBColor(67,67,67), 16, 4)
]:
    st = styles[name]; st.font.name = "Arial"; st.font.size = Pt(size); st.font.color.rgb = color
    st.font.bold = False; st._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

# Plain title paragraph: required for Google Docs-targeted output.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Local Network Messenger in Go")
set_font(r, size=26, bold=False)
add_p(doc, "A beginner mini tutorial for macOS and local Wi-Fi networks", after=4, size=13, color=GRAY)
add_p(doc, "Build a terminal chat app first. Add file transfer next. Wrap it in a simple desktop interface last.", after=14, size=11)
add_note(doc, "Learning approach", "Work in small steps. At the end of each milestone, run the app and prove that one new behavior works before moving on.")

doc.add_heading("What You Will Build", level=1)
add_p(doc, "The first version is intentionally small: one Go server listens for TCP connections and one Go client connects to it. Two computers on the same Wi-Fi network can exchange text messages and send files. The terminal version teaches the networking basics before a GUI adds extra moving parts.")
for t in ["Milestone 1: install Go and create the project.", "Milestone 2: run a TCP server.", "Milestone 3: connect a client and exchange text.", "Milestone 4: test across your local network.", "Milestone 5: add a small file-transfer protocol.", "Milestone 6: add an optional Fyne desktop window."]:
    add_bullet(doc, t)
add_note(doc, "Keep v1 local", "Do not expose this app directly to the internet. Encryption, authentication, peer discovery, and relay servers belong in a later project.")

doc.add_heading("1. Set Up Go on macOS", level=1)
add_step(doc, "Download the macOS installer from https://go.dev/doc/install and follow the package prompts.")
add_step(doc, "Close and reopen Terminal, then verify the installation:")
add_code(doc, "go version")
add_step(doc, "Create a new project folder and initialize a Go module:")
add_code(doc, """mkdir lion-chat
cd lion-chat
go mod init example.com/lion-chat
mkdir -p cmd/server cmd/client internal/protocol downloads""")
add_p(doc, "Your starting layout should look like this:", after=4)
add_code(doc, """lion-chat/
  cmd/server/main.go
  cmd/client/main.go
  internal/protocol/
  downloads/
  go.mod""")
add_note(doc, "Checkpoint", "Run go version and confirm that go.mod exists inside lion-chat.")

page_break(doc)
doc.add_heading("2. Build the TCP Server", level=1)
add_p(doc, "A TCP server needs three basic actions: listen on a port, accept a connection, and read from that connection. The Go standard library provides these through the net package.")
add_p(doc, "Create cmd/server/main.go:", after=4, bold=True)
add_code(doc, r'''package main

import (
    "bufio"
    "fmt"
    "log"
    "net"
)

func main() {
    listener, err := net.Listen("tcp", ":9000")
    if err != nil {
        log.Fatal(err)
    }
    defer listener.Close()
    log.Println("server listening on :9000")

    for {
        conn, err := listener.Accept()
        if err != nil {
            log.Println("accept error:", err)
            continue
        }
        go handleConnection(conn)
    }
}

func handleConnection(conn net.Conn) {
    defer conn.Close()
    log.Println("connected:", conn.RemoteAddr())

    scanner := bufio.NewScanner(conn)
    for scanner.Scan() {
        fmt.Println("message:", scanner.Text())
    }
    if err := scanner.Err(); err != nil {
        log.Println("read error:", err)
    }
}''')
add_p(doc, "Run the server:", after=4)
add_code(doc, "go run ./cmd/server")
add_note(doc, "What to notice", "The accept loop keeps listening. Each connection is handled in a goroutine, so a later version can support more than one client.")

doc.add_heading("3. Build the TCP Client", level=1)
add_p(doc, "The client connects with net.Dial, reads lines typed in Terminal, and writes each line to the TCP connection.")
add_p(doc, "Create cmd/client/main.go:", after=4, bold=True)
add_code(doc, r'''package main

import (
    "bufio"
    "fmt"
    "log"
    "net"
    "os"
)

func main() {
    conn, err := net.Dial("tcp", "localhost:9000")
    if err != nil {
        log.Fatal(err)
    }
    defer conn.Close()

    fmt.Println("connected; type a message and press Return")
    input := bufio.NewScanner(os.Stdin)
    for input.Scan() {
        if _, err := fmt.Fprintln(conn, input.Text()); err != nil {
            log.Fatal(err)
        }
    }
}''')
add_p(doc, "Open a second Terminal window and run:", after=4)
add_code(doc, "go run ./cmd/client")
add_note(doc, "Checkpoint", "Type hello in the client Terminal. The server Terminal should print message: hello.")

page_break(doc)
doc.add_heading("4. Turn It Into a Two-Way Chat", level=1)
add_p(doc, "The first client only sends. Add a goroutine so the client can also print messages received from the server.")
add_p(doc, "In cmd/client/main.go, start this goroutine after connecting:", after=4, bold=True)
add_code(doc, r'''go func() {
    scanner := bufio.NewScanner(conn)
    for scanner.Scan() {
        fmt.Println("\npeer:", scanner.Text())
        fmt.Print("> ")
    }
}()''')
add_p(doc, "For a simple echo test, make the server write each line back to the connected client. Add this inside the server scan loop:", after=4, bold=True)
add_code(doc, r'''message := scanner.Text()
fmt.Println("message:", message)
if _, err := fmt.Fprintln(conn, message); err != nil {
    log.Println("write error:", err)
    return
}''')
add_note(doc, "Checkpoint", "Send a line from the client. You should see it in the server Terminal and echoed back in the client Terminal.")
add_p(doc, "A real two-person chat needs either two clients plus server-side broadcasting, or a direct peer-to-peer design. For a beginner project, keep the server as the hub: store connected clients and broadcast each text message to the others.")

doc.add_heading("5. Test Across Your Local Network", level=1)
add_p(doc, "Once localhost works, run the server on one Mac and connect from another computer on the same Wi-Fi network.")
add_step(doc, "On the server Mac, find the local Wi-Fi address:")
add_code(doc, "ipconfig getifaddr en0")
add_step(doc, "If that prints an address such as 192.168.1.20, replace localhost in the client:")
add_code(doc, r'''conn, err := net.Dial("tcp", "192.168.1.20:9000")''')
add_step(doc, "Run the server first, then run the client from the second computer.")
add_note(doc, "macOS firewall", "macOS may ask whether Terminal or your compiled app should accept incoming connections. Allow it for local testing. Both devices must be on the same network.")

doc.add_heading("Quick Troubleshooting", level=2)
for t in ["Connection refused: start the server first and confirm the port is 9000.", "Timeout: confirm both devices are on the same Wi-Fi network and recheck the server IP address.", "No IP from en0: your active adapter may be different. Run ifconfig and look for a private address such as 192.168.x.x or 10.x.x.x.", "Messages stick together: keep using newline-delimited text until the framed protocol is added."]:
    add_bullet(doc, t)

page_break(doc)
doc.add_heading("6. Add a Small Framed Protocol", level=1)
add_p(doc, "Text lines are enough for the first chat test, but file transfer needs structure. Use a small header before every payload so the receiver knows what is arriving and how many bytes to read.")
add_p(doc, "A simple teaching protocol:", after=4)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = False
headers = ["Type", "Purpose", "Payload"]
for i, h in enumerate(headers): set_cell_text(table.rows[0].cells[i], h, bold=True)
mark_header(table.rows[0])
for row in [
    ("TEXT", "Send one chat message", "UTF-8 text bytes"),
    ("FILE_META", "Describe an incoming file", "Filename and byte count"),
    ("FILE_DATA", "Stream file contents", "Raw file bytes"),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row): set_cell_text(cells[i], value)
add_p(doc, "", after=2)
add_p(doc, "Keep the wire format simple: a one-line header followed by exactly the announced number of payload bytes.")
add_code(doc, r'''TEXT 12
hello there!

FILE_META 18
photo.jpg|245760

FILE_DATA 245760
<raw bytes follow>''')
add_note(doc, "Important", "Do not use Scanner for raw file bytes. Parse the short header line, then use io.CopyN to read the exact payload size.")

doc.add_heading("7. Implement File Transfer", level=1)
add_p(doc, "Start with one file at a time. The sender opens a file, sends metadata, then streams the file bytes. The receiver validates the filename and size before writing into downloads/.")
add_p(doc, "Sender outline:", after=4, bold=True)
add_code(doc, r'''file, err := os.Open(path)
if err != nil { return err }
defer file.Close()

info, err := file.Stat()
if err != nil { return err }

name := filepath.Base(info.Name())
fmt.Fprintf(conn, "FILE_META %d\n%s|%d", len(name)+1+len(strconv.FormatInt(info.Size(), 10)), name, info.Size())
fmt.Fprintf(conn, "FILE_DATA %d\n", info.Size())
_, err = io.Copy(conn, file)
return err''')
add_p(doc, "Receiver outline:", after=4, bold=True)
add_code(doc, r'''safeName := filepath.Base(receivedName)
if safeName != receivedName || safeName == "." {
    return errors.New("unsafe filename")
}
if size < 0 || size > 50<<20 {
    return errors.New("file too large")
}

dst := filepath.Join("downloads", safeName)
if _, err := os.Stat(dst); err == nil {
    return errors.New("file already exists")
}
out, err := os.Create(dst)
if err != nil { return err }
defer out.Close()

_, err = io.CopyN(out, conn, size)
return err''')
add_note(doc, "Guardrails", "Reject unsafe filenames, cap file size, handle disconnects, create downloads/ if needed, and never overwrite an existing file silently.")

page_break(doc)
doc.add_heading("8. Organize the Code", level=1)
add_p(doc, "Once file transfer works, move shared framing code into internal/protocol. Keep command entry points small.")
add_code(doc, """lion-chat/
  cmd/server/main.go       # listen, accept, broadcast
  cmd/client/main.go       # connect, terminal input, receive loop
  internal/protocol/
    message.go             # frame types and parsing
    file.go                # send and receive helpers
  downloads/               # received files
  go.mod""")
add_p(doc, "Useful protocol API shape:", after=4)
add_code(doc, r'''type Frame struct {
    Type    string
    Size    int64
    Payload io.Reader
}

func WriteText(w io.Writer, message string) error
func SendFile(w io.Writer, path string) error
func ReadFrame(r *bufio.Reader) (Frame, error)
func ReceiveFile(r io.Reader, name string, size int64) error''')
add_note(doc, "Exercise", "Write a table-driven test for unsafe filenames and files larger than your chosen limit.")

doc.add_heading("9. Optional Desktop GUI", level=1)
add_p(doc, "Add a GUI only after the terminal version works. That keeps networking bugs separate from interface bugs.")
doc.add_heading("Recommended: Fyne", level=2)
add_p(doc, "Fyne is the simplest fit for this project because it stays in Go and supports desktop apps, including macOS. Start with a single window containing:")
for t in ["Server address field", "Connect button", "Scrollable message history", "Message input and Send button", "Choose File button", "Connection status label"]:
    add_bullet(doc, t)
add_p(doc, "Your UI callbacks should call the same protocol helpers used by the terminal client. Run network reads in a goroutine and update the message view when a frame arrives.")
add_code(doc, """go get fyne.io/fyne/v2@latest

# Then follow the official getting-started guide:
# https://docs.fyne.io/started/""")

doc.add_heading("Fyne vs. Wails", level=2)
table = doc.add_table(rows=1, cols=3); table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, h in enumerate(["Choice", "Use it when", "Tradeoff"]): set_cell_text(table.rows[0].cells[i], h, bold=True)
mark_header(table.rows[0])
for row in [
    ("Fyne", "You want the most direct beginner path.", "Pure Go and simple; visual customization is more limited."),
    ("Wails", "You already know web UI tools or want a polished custom interface.", "Adds HTML, CSS, JavaScript, Node, and platform setup."),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row): set_cell_text(cells[i], value)
add_p(doc, "", after=2)
add_note(doc, "Recommendation", "Build the first desktop version with Fyne. Revisit Wails only after the protocol layer is stable.")

page_break(doc)
doc.add_heading("10. Finish Line Checklist", level=1)
for t in [
    "Go is installed and go version works.",
    "The server starts on :9000.",
    "A local client connects through localhost.",
    "A second computer connects through the server Mac's local IP address.",
    "Text messages travel in both directions.",
    "Files arrive in downloads/ with size and filename checks.",
    "Existing files are not overwritten silently.",
    "The optional GUI reuses the protocol package instead of duplicating networking logic.",
]:
    add_bullet(doc, t)

doc.add_heading("Manageable Stretch Goals", level=1)
for t in ["Add usernames and timestamps.", "Broadcast messages to multiple connected clients.", "Show file-transfer progress.", "Save a small local message history.", "Add tests for malformed headers and interrupted transfers."]:
    add_bullet(doc, t)

doc.add_heading("Keep These Out of v1", level=1)
for t in ["Internet exposure", "Encryption and identity verification", "Automatic peer discovery", "Cloud relay servers", "Account systems"]:
    add_bullet(doc, t)
add_p(doc, "Those are worthwhile follow-up projects, but each changes the threat model and design. Finish the local-network version first.")

doc.add_heading("References", level=1)
for t in [
    "Install Go: https://go.dev/doc/install",
    "Go net package: https://pkg.go.dev/net",
    "Fyne documentation: https://docs.fyne.io/",
    "Wails installation guide: https://wails.io/docs/gettingstarted/installation",
    "Optional practice: Boot.dev Go learning paths can provide extra exercises. Use official documentation as the source of truth for APIs and installation.",
]:
    add_bullet(doc, t)

# Footer only; Google Docs preset avoids decorative running furniture.
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Local Network Messenger in Go  |  Beginner Mini Tutorial")
set_font(r, size=9, color=GRAY)

doc.core_properties.title = "Local Network Messenger in Go"
doc.core_properties.subject = "Beginner mini tutorial for a Go TCP messaging app with file transfer"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
